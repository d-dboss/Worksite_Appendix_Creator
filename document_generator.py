"""
Photo Appendix Generator - Document Generation Module
This module handles creating the output document with photos and metadata.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL # For vertical alignment in cells

# Import cleanup functions with specific names to avoid confusion
from photo_processor import cleanup_temp_files as cleanup_photo_temp_files
from map_generator import generate_map_image, generate_compass_indicator, cleanup_temp_files as cleanup_map_temp_files

# Removed the problematic set_cell_margins function

def create_document(photo_data_list, output_path, images_per_page=2, include_location=True):
    """
    Create a Word document with photos, captions, location maps, and compass indicators.

    Args:
        photo_data_list (list): List of dictionaries containing photo data.
        output_path (str): Path to save the output document.
        images_per_page (int): Number of images per page (1, 2, or 4).
        include_location (bool): Whether to include location data, map, and compass.

    Returns:
        bool: True if successful, False otherwise.
    """
    print("\n--- Starting Document Generation ---")
    print(f"Output Path: {output_path}")
    print(f"Images per page: {images_per_page}")
    print(f"Include Location Info: {include_location}")

    temp_map_compass_files = [] # Keep track of map/compass files created *during this run*

    try:
        # Create a new document
        doc = Document()

        # Set document properties (optional but good practice)
        doc.core_properties.title = "Photo Appendix"
        doc.core_properties.author = "Photo Appendix Generator"

        # --- Page Setup ---
        # Assuming standard Letter size (8.5x11 inches) with 1-inch margins
        # Available content width = 8.5 - 1 (left margin) - 1 (right margin) = 6.5 inches
        page_content_width = Inches(6.5)
        print(f"Effective Page Content Width: {page_content_width / Inches(1):.2f} inches")

        # Set margins for all sections (usually just one)
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            # Optional: Header/Footer distance
            # section.header_distance = Inches(0.5)
            # section.footer_distance = Inches(0.5)


        # --- Document Title ---
        title = doc.add_heading("Photo Appendix", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add space after title
        title.paragraph_format.space_after = Pt(12)


        # --- Introduction Text ---
        intro_text = ("This document contains photographs and associated metadata. "
                      "Where available and selected, location information including "
                      "a map marker and compass direction (camera heading) is provided below each photo.")
        p_intro = doc.add_paragraph(intro_text)
        p_intro.paragraph_format.space_after = Pt(18) # Add space after intro


        # --- Photo Entries ---
        total_photos = len(photo_data_list)
        for i, photo_data in enumerate(photo_data_list):
            print(f"\nAdding Photo {i+1} of {total_photos}: {photo_data['filename']}")

            # --- Page Breaks ---
            # Add page break *before* the item if it's not the first item AND
            # it's the start of a new page based on images_per_page setting.
            if i > 0 and i % images_per_page == 0:
                 doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                 print("  Added Page Break.")

            # --- Image Insertion ---
            # Use the temporary converted file if it exists, otherwise use original path
            image_path = photo_data.get('temp_file') or photo_data['path']
            if not image_path or not os.path.exists(image_path):
                 print(f"  ERROR: Image file not found for {photo_data['filename']} at path: {image_path}")
                 # Add an error note in the document
                 error_para = doc.add_paragraph()
                 error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                 error_run = error_para.add_run(f"[Error: Image file not found for {photo_data['filename']}]")
                 error_run.font.italic = True
                 error_run.font.color.rgb = RGBColor(255, 0, 0)
                 continue # Skip to the next photo


            # Determine image width based on layout
            # Simple stacking layout assumed for 1, 2.
            # For 4, we assume 2x2 grid - needs table for side-by-side. Let's stick to stacking for now.
            # TODO: Implement 2x2 grid layout using tables if needed.
            if images_per_page == 1:
                img_width = page_content_width
                loc_table_width = page_content_width # Location info spans full width
            elif images_per_page == 2:
                img_width = page_content_width # Stacked vertically, full width
                loc_table_width = page_content_width
            elif images_per_page >= 4: # Treat 4+ as half width, stacked
                img_width = (page_content_width - Inches(0.2)) / 2 # Allow small gap if side-by-side
                loc_table_width = img_width # Location info below each half-width image
                print("  Using half-width layout for 4+ images per page (assuming stacked).")
            else: # Default to full width
                 img_width = page_content_width
                 loc_table_width = page_content_width

            try:
                # Add Picture - Centered
                # Create a paragraph, center it, then add the run with the picture
                img_paragraph = doc.add_paragraph()
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_paragraph.add_run()
                # Check if dimensions are available to avoid stretching small images too much
                img_height = None
                # Use oriented dimensions if available
                p_width = photo_data.get('width')
                p_height = photo_data.get('height')
                if p_width and p_height and p_width > 0: # Check width > 0
                     aspect_ratio = p_height / p_width
                     img_height = img_width * aspect_ratio
                     # Optional: Limit max height? e.g., max_height = Inches(5)
                     # if img_height > max_height:
                     #      img_height = max_height
                     #      img_width = img_height / aspect_ratio

                if img_height is not None and img_height > 0:
                     img_run.add_picture(image_path, width=img_width, height=img_height)
                else:
                     img_run.add_picture(image_path, width=img_width) # Let Word determine height

                # Reduce space after image paragraph if needed
                img_paragraph.paragraph_format.space_after = Pt(2)


            except Exception as e:
                print(f"  ERROR: Failed to add image {photo_data['filename']} to document: {str(e)}")
                # Add error placeholder in document
                error_para = doc.add_paragraph()
                error_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                error_run = error_para.add_run(f"[Error adding image: {photo_data['filename']}]")
                error_run.font.italic = True
                error_run.font.color.rgb = RGBColor(255, 0, 0)
                # Optionally continue to add caption/location info if desired
                # continue


            # --- Caption Insertion ---
            caption_paragraph = doc.add_paragraph()
            caption_text = photo_data.get('caption', '') or "No caption available"
            caption_run = caption_paragraph.add_run(caption_text)
            caption_run.font.size = Pt(9) # Slightly smaller caption font
            caption_run.font.italic = True # Italicize caption
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_paragraph.paragraph_format.space_after = Pt(6) # Space after caption


            # --- Location Information Section (Map, Coordinates, Compass) ---
            if include_location:
                # Check if *any* relevant data exists
                has_gps = photo_data.get('latitude') is not None and photo_data.get('longitude') is not None
                has_orientation = photo_data.get('orientation') is not None

                if has_gps or has_orientation:
                    print(f"  Adding location section (GPS: {has_gps}, Orientation: {has_orientation})")

                    # --- Create Location Table ---
                    # Using a 2-column table: [Map + Coords] | [Compass]
                    location_table = doc.add_table(rows=1, cols=2)
                    location_table.style = 'Table Grid' # Add borders for visual structure
                    location_table.autofit = False
                    location_table.allow_autofit = False # Prevent auto resizing

                    # Set column widths (adjust proportions as needed)
                    # Approx 60% for Map/Coords, 40% for Compass seems reasonable
                    col1_width = loc_table_width * 0.60
                    col2_width = loc_table_width * 0.40
                    location_table.columns[0].width = int(col1_width)
                    location_table.columns[1].width = int(col2_width)

                    # Get cell objects
                    map_coords_cell = location_table.cell(0, 0)
                    compass_cell = location_table.cell(0, 1)

                    # Set vertical alignment for cells
                    map_coords_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    compass_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # ***** FIX: Removed call to set_cell_margins *****
                    # Instead, we will set spacing on paragraphs inside the cells.


                    # --- Populate Left Cell (Map + Coordinates) ---
                    # Ensure the cell has at least one paragraph to start with
                    if not map_coords_cell.paragraphs: map_coords_cell.add_paragraph()
                    # Clear default paragraph content if any (usually just empty)
                    map_coords_cell.paragraphs[0].text = ''
                    for para in map_coords_cell.paragraphs[1:]: # Remove extra default paragraphs if created somehow
                         map_coords_cell._element.remove(para._element)

                    map_path = None
                    if has_gps:
                        print(f"    Generating map for Lat={photo_data['latitude']:.5f}, Lon={photo_data['longitude']:.5f}")
                        map_path = generate_map_image(
                            photo_data['latitude'],
                            photo_data['longitude'],
                            zoom=15,
                            size=(180, 180) # Slightly larger map image size
                        )
                        if map_path:
                            temp_map_compass_files.append(map_path)
                            print(f"    Map generated: {map_path}")
                        else:
                            print("    Map generation failed.")

                    # Add Map Image (if generated) to the first paragraph
                    map_para = map_coords_cell.paragraphs[0]
                    map_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add spacing for padding effect
                    map_para.paragraph_format.space_before = Pt(3)
                    map_para.paragraph_format.space_after = Pt(2) # Small space after map

                    if map_path:
                        try:
                            map_run = map_para.add_run()
                             # Adjust display size in document (e.g., 1.75 inches wide)
                            map_display_width = Inches(1.75)
                            map_run.add_picture(map_path, width=map_display_width)
                        except Exception as e:
                            print(f"    ERROR: Could not add map image {map_path} to document: {e}")
                            # Add error text in a new paragraph if map failed
                            err_p = map_coords_cell.add_paragraph("(Map Error)")
                            err_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif has_gps: # GPS existed, but map failed
                        map_para.add_run("(Map Unavailable)").italic = True


                    # Add Coordinates Text (if GPS available) in a new paragraph
                    if has_gps:
                         coords_para = map_coords_cell.add_paragraph()
                         coords_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                         coords_text = f"Lat: {photo_data['latitude']:.5f}\nLon: {photo_data['longitude']:.5f}" # Use newline
                         coords_run = coords_para.add_run(coords_text)
                         coords_run.font.size = Pt(8) # Smaller font for coordinates
                         # Add spacing for padding effect
                         coords_para.paragraph_format.space_before = Pt(2)
                         coords_para.paragraph_format.space_after = Pt(3)
                    elif not map_path: # No GPS and no map image was attempted/added
                         # Add placeholder text if the cell would otherwise be empty
                         no_gps_para = map_coords_cell.paragraphs[0] # Use first para if map wasn't added
                         no_gps_para.add_run("(No GPS Data)").italic = True
                         no_gps_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


                    # --- Populate Right Cell (Compass) ---
                    # Ensure the cell has at least one paragraph to start with
                    if not compass_cell.paragraphs: compass_cell.add_paragraph()
                     # Clear default paragraph content if any (usually just empty)
                    compass_cell.paragraphs[0].text = ''
                    for para in compass_cell.paragraphs[1:]:
                         compass_cell._element.remove(para._element)

                    compass_path = None
                    if has_orientation:
                         orientation = photo_data['orientation']
                         print(f"    Generating compass for Orientation={orientation:.1f}°")
                         compass_path = generate_compass_indicator(orientation, size=(100, 100)) # Keep compass size reasonable

                         if compass_path:
                              temp_map_compass_files.append(compass_path)
                              print(f"    Compass generated: {compass_path}")
                         else:
                              print("    Compass generation failed.")


                    # Add Compass Image (if generated) to the first paragraph
                    compass_para = compass_cell.paragraphs[0]
                    compass_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Add spacing for padding effect
                    compass_para.paragraph_format.space_before = Pt(3)
                    compass_para.paragraph_format.space_after = Pt(3)

                    if compass_path:
                        try:
                             comp_run = compass_para.add_run()
                             # Adjust display size (e.g., 1 inch wide)
                             comp_display_width = Inches(1.0)
                             comp_run.add_picture(compass_path, width=comp_display_width)
                        except Exception as e:
                             print(f"    ERROR: Could not add compass image {compass_path} to document: {e}")
                             err_p = compass_cell.add_paragraph("(Compass Error)")
                             err_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif has_orientation: # Orientation existed, but compass failed
                         # Display orientation as text if compass image failed
                         comp_text_para = compass_cell.paragraphs[0] # Use first para
                         comp_text_run = comp_text_para.add_run(f"Direction:\n{orientation:.1f}°")
                         comp_text_run.font.size = Pt(8)
                    else: # No orientation data
                         no_orient_para = compass_cell.paragraphs[0] # Use first para
                         no_orient_para.add_run("(No Direction Data)").italic = True


                    # Add space after the location table
                    # Use paragraph spacing on the *last* paragraph added *after* the table
                    # (or add an empty paragraph with spacing)
                    spacer_para = doc.add_paragraph()
                    spacer_para.paragraph_format.space_before = Pt(6) # Add space below table


                # else: (No GPS and No Orientation) - No location table added


            # --- Spacing Between Photos on Same Page ---
            # Add space *before* the next photo if it's on the same page
            # Check if the *next* item (i+1) does NOT start a new page
            if (i + 1) < total_photos and (i + 1) % images_per_page != 0:
                 # Add visual separator or just space
                 # doc.add_paragraph("----" * 10).alignment = WD_ALIGN_PARAGRAPH.CENTER # Visual separator
                 spacer_para = doc.add_paragraph() # Just add space
                 spacer_para.paragraph_format.space_before = Pt(12)
                 spacer_para.paragraph_format.space_after = Pt(6)


        # --- End of Loop ---

        # Save the document
        print(f"\nSaving document to: {output_path}")
        doc.save(output_path)
        print("Document saved successfully.")
        return True

    except Exception as e:
        print(f"\n--- FATAL ERROR DURING DOCUMENT GENERATION ---")
        print(f"Error Type: {type(e).__name__}")
        print(f"Error Details: {str(e)}")
        import traceback
        print("--- Traceback ---")
        traceback.print_exc()
        print("-----------------")
        # Attempt to save even on error? Maybe not useful.
        return False

    finally:
        # --- Cleanup ---
        # Regardless of success or failure, clean up temporary files
        print("\nPerforming final cleanup...")
        # Cleanup temporary files created by photo_processor (e.g., HEIC conversions)
        cleanup_photo_temp_files(photo_data_list)
        # Cleanup temporary files created by map_generator (maps, compasses)
        cleanup_map_temp_files(temp_map_compass_files) # Pass the list of files created in *this run*
        print("Cleanup finished.")
        print("--- Document Generation Process Ended ---")